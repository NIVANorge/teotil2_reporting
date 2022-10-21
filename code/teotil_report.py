import os

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

plt.style.use("ggplot")


def insert_para_after(para, style="Normal", align="center"):
    """Insert a new paragraph after the given paragraph.

    Args:
        para:     Obj. Paragraph object to insert after
        style:    Str. One of Word's pre-defined styles
        align:    Str. One of 'left', 'center' or 'right'

    Returns:
        Paragraph object.
    """
    new_p = OxmlElement("w:p")
    para._p.addnext(new_p)
    new_para = Paragraph(new_p, para._parent)

    new_para.style = style

    if align == "center":
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "right":
        new_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        raise ValueError("'align' must be one of 'left', 'center' or 'right'.")

    return new_para


def insert_image(para, img_path, width_cm=15):
    """Insert an image into a paragraph.

    Args:
        para:     Obj. Paragraph object to insert after
        img_path: Str. Path to image
        width_cm: Float. Width of image in cm

    Returns:
        Run object.
    """
    run = para.add_run()
    run.add_picture(img_path, width=Cm(width_cm))

    return run


def insert_table(doc, para, df, table_style="Grid Table 4 Accent 1"):
    """Insert a dataframe as a table into the given paragraph.

    Args:
        doc:         Obj. Document object
        para:        Obj. Paragraph object to insert after
        df:          Obj. Dataframe with table data
        table_style: Str. Word table style to apply. NOTE: If your template document
                     doesn't contain any styled tables, the list of available choices
                     may be very limited. Try

                         styles = doc.styles
                         table_styles = [s for s in styles if s.type == WD_STYLE_TYPE.TABLE]

                     to see a list of available styles in your document. To add a style,
                     manually create a table with the desired style in the template, then save
                     it, delete the table and save again. The chosen style should now be
                     available

    Returns:
        None.
    """
    df = df.copy().astype(str)
    df.replace("<NA>", "", inplace=True)

    # Tables can only be added at the end of the document, but they can be moved
    # elsewhere afterwards. See
    # https://github.com/python-openxml/python-docx/issues/156
    table = doc.add_table(df.shape[0] + 1, df.shape[1])
    table.style = table_style

    # Add header
    for j in range(df.shape[-1]):
        table.cell(0, j).text = df.columns[j]
        table.cell(0, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add data
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            table.cell(i + 1, j).text = str(df.values[i, j])
            table.cell(i + 1, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Move to desired location (paragraph)
    tbl, p = table._tbl, para._p
    p.addnext(tbl)


def make_plot(df, title, base_path, year):
    """Create a time series plot from a dataframe.

    Args:
        df:        Obj. Dataframe with table data
        title:     Str. Title for plot
        base_path: Str. File path for original CSV data table

    Returns:
        Plot path. Plot is saved to ./plots/
    """
    df = df.copy()

    if base_path[-5] == "p":
        ylabel = "Fosfor (tonn)"
    elif base_path[-5] == "n":
        ylabel = "Nitrogen (tonn)"
    else:
        raise ValueError("Could not identify ylabel.")

    ax = df.set_index("År").plot(figsize=(8, 6), title=title, xlabel="", fontsize=12)
    ax.legend(
        loc="center", bbox_to_anchor=(0.5, -0.2), ncol=3, prop={"size": 12}
    ).get_frame().set_boxstyle("Round", pad=0.2, rounding_size=0.5)
    ax.set_ylabel(ylabel, fontdict={"fontsize": 12, "fontweight": "bold"})
    plt.tight_layout()

    fname = os.path.split(base_path)[1][:-4]
    png_path = f"../report_{year}/plots/{fname}_plot.png"
    plt.savefig(png_path, dpi=200)
    plt.close()

    return png_path


def read_data_table(fpath, heading):
    """Reads a CSV to a dataframe. Rounds values to integers and adds
    an extra row for 1985 at the start of each time series.

    Args:
        fpath: Path to CSV data

    Returns:
        Dataframe
    """
    # Read 1985 data
    csv_path = r"../jse_data_1985/data_1985.csv"
    df_85 = pd.read_csv(csv_path, index_col="section_name")

    # Read model data
    df = pd.read_csv(fpath)

    # Correct typo in Jose's column names
    df.rename({"Bakgrun": "Bakgrunn"}, axis="columns", inplace=True)

    # Insert data from 1985 if relevant
    if heading in df_85.index:
        row_85 = df_85.loc[[heading]]
        df = pd.concat([row_85, df], axis="rows")

    df.reset_index(inplace=True, drop=True)
    df.sort_values("År", inplace=True)

    # Rename col as requested by Rita Vigdis
    if "Befolkning" in df.columns:
        df.rename({"Befolkning": "Avløp"}, axis="columns", inplace=True)

    header = [
        "År",
        "Akvakultur",
        "Jordbruk",
        "Avløp",
        "Industri",
        "Bakgrunn",
        "Totalt",
        "Menneskeskapt",
    ]
    df = df[header]
    df = df.round(0).astype("Int64")

    return df


def patch_legacy_data(df, legacy_data_fold, fname, heading, cutoff_year):
    """The data in NIVA's database does not exactly match what is reported in old TEOTIL
    reports. Differences are especially obvious in early years (pre-1995). It looks as
    though the data used in previous reports comes from Access databases where John Rune
    manually patched gaps in the data series from the database. In most cases the values
    just seem to be averaged/filled from neighbouring years. I am not able to reproduce
    these values using data in our database, but for consistency it is useful to be able
    to use the old values for early years in new reports. This function takes the 'df'
    generated from data in the database and patches it with the manually edited legacy
    data used by John Rune. Legacy data is used for years less than or equal to the
    'cutoff_year', and data from the database is used after.

    Args
        df:               Dataframe. Original, 'unmodified' data from NIVA's database
        legacy_data_fold: Str. Folder containing legacy data files used by Jose and John
                          Rune
        fname:            Str. CSV to process
        heading:          Str. Document heading to process
        cutoff_year:      Int. Legacy data will be used for years less than or equal to
                          this

    Returns
        Datafarme.
    """
    fpath = os.path.join(legacy_data_fold, fname)
    legacy_df = read_data_table(fpath, heading)

    legacy_df = legacy_df[legacy_df["År"] <= cutoff_year]
    df = df[df["År"] > cutoff_year]
    df = pd.concat([legacy_df, df], axis="rows")
    df.sort_values("År", inplace=True)

    return df


def get_teotil_results_main_catchments(st_yr, end_yr):
    """ """
    # List of catchments flowing to coast. 315 flows into Skagerrak
    main_catches = [f"{i:03d}." for i in range(1, 248)] + ["315."]
    df_list = []
    for year in range(st_yr, end_yr + 1):
        base_url = f"https://raw.githubusercontent.com/NIVANorge/teotil2/main/data/norway_annual_output_data/teotil2_results_{year}.csv"
        df = pd.read_csv(base_url)
        df = df.query("regine in @main_catches").copy()
        df["År"] = year
        cols = [i for i in df.columns if i.split("_")[0] == "accum"]
        df = df[["regine", "År"] + cols]
        df_list.append(df)
    df = pd.concat(df_list)

    return df


def get_aggregation_dict_for_columns(par):
    """Make a dict mapping TEOTIL column names to columns used in the report
       with aggregation where necessary.

    Args
        par: Str. Either 'n' or 'p'

    Returns
        Dict with key's equal to headings used in the report and values are lists
        of columns to aggregate in the TEOTIL output.
    """
    assert par in ("n", "p")

    agg_dict = {
        "Akvakultur": [f"accum_aqu_tot-{par}_tonnes"],
        "Jordbruk": [
            f"accum_agri_diff_tot-{par}_tonnes",
            f"accum_agri_pt_tot-{par}_tonnes",
        ],
        "Avløp": [f"accum_ren_tot-{par}_tonnes", f"accum_spr_tot-{par}_tonnes"],
        "Industri": [f"accum_ind_tot-{par}_tonnes"],
        "Bakgrunn": [
            f"accum_nat_diff_tot-{par}_tonnes",
            f"accum_urban_tot-{par}_tonnes",
        ],
        "Totalt": [f"accum_all_sources_tot-{par}_tonnes"],
        "Menneskeskapt": [
            f"accum_anth_diff_tot-{par}_tonnes",
            f"accum_all_point_tot-{par}_tonnes",
        ],
        # "Urban": [
        #     f"accum_urban_tot-{par}_tonnes",
        # ],
    }

    return agg_dict


def aggregate_parameters(df, par):
    """Aggregate columns in TEOTIL output to headings used in the report.

    Args
        df:  Dataframe of TEOTIL results
        par: Str. Either 'n' or 'p'

    Returns
        Dataframe.
    """
    agg_dict = get_aggregation_dict_for_columns(par)
    for group, cols in agg_dict.items():
        df[group] = df[cols].sum(axis=1)

    df = df[["regine", "År"] + list(agg_dict.keys())]

    return df


def aggregate_regions(df, par, out_fold=None):
    """Sum TEOTIL output for the main catchments for each region defined in the
       report.

    Args
        df:       Dataframe of results aggregated to the correct column anmes for
                  the report
        par:      Str. Either 'n' or 'p'
        out_fold: Bool or str. Default None. Folder to save CSVs to, if desired

    Returns
        Dict of dataframes. Optionall, results for each region are saved to CSV.
    """
    assert par in ("n", "p")

    # Map regions used in report to main catchments
    # Intervals are "Python-style" i.e. include first but not last element in range
    regions_dict = {
        # Definerte kystavsnitt (chapter 5)
        "Norges kystområder": [1, 248, 315],
        "Sverige – Strømtangen fyr": [1, 3],
        "Indre Oslofjord": [5, 10],
        "Svenskegrensa – Lindesnes": [1, 24],
        "Lindesnes – Stad": [24, 92],
        "Stad – Russland": [92, 248],
        # Norske vannregioner (chapter 6)
        "Glomma": [1, 11],
        "Vest-Viken": [11, 18],
        "Agder": [18, 27],
        "Rogaland": [27, 41],
        "Hordaland": [41, 68],
        "Sogn og Fjordane": [68, 92],
        "Møre og Romsdal": [92, 117],
        "Trøndelag": [117, 144],
        "Nordland": [144, 186],
        "Troms": [186, 211],
        "Finnmark": [211, 248],
        # Norske forvaltingsplanområder (chapter 7)
        "Nordsjøen": [1, 91, 315],  # 315 is included here too in John Rune's Access db
        "Norskehavet": [91, 171],
        "Barentshavet": [171, 248],
    }

    result_dict = {}
    for region, catches in regions_dict.items():
        if len(catches) == 2:
            catch_list = list(range(catches[0], catches[1]))
        else:
            catch_list = list(range(catches[0], catches[1])) + [catches[2]]
        catch_list = [f"{i:03d}." for i in catch_list]

        reg_df = df.query("regine in @catch_list").copy()
        reg_df = reg_df.groupby("År").sum().reset_index()
        reg_df = reg_df.round(0).astype(int)
        result_dict[region] = reg_df

        if out_fold:
            csv_path = os.path.join(out_fold, f"{region}_{par}.csv")
            reg_df.to_csv(csv_path, index=False)

    return result_dict


def filename_from_heading(heading):
    """Build a file name based on chapter headings in the Word template."""
    name, par = heading.split(":")
    if par[1] == "f":
        par = "p"
    else:
        par = "n"

    return f"{name}_{par}.csv"
